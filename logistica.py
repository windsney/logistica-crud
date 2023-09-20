from tkinter import *
from tkinter import ttk
from tkinter import messagebox
import  tkinter as tk
from docx import Document
import customtkinter
import sqlite3
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

janela = Tk()
p1 = PhotoImage(file='icone_caixa.png')
janela.iconphoto(False, p1)

class Relatorios():
    def printProduto(self):
       pass

    def geraRelatorioProduto(self):


        


        banco = sqlite3.connect('produtos.bd')
        cursor = banco.cursor()
        cursor.execute('SELECT cod,nome_produto, quantidade, categoria FROM produtos')
        dados1 = cursor.fetchall()
        doc = Document()

        # Adicionar um título
        titlo=doc.add_heading('Estoque 2 Irmãos', level=1)
        titlo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tm=titlo.runs[0]
        tm.font.size = Pt(26)
        tm.font.name = 'Arial'
        for dado in dados1:
            cod=dado[0]
            nome=dado[1]
            qtd=dado[2]
            cate=dado[3]


            self.codigoRel = cod
            self.produtoRel = nome
            self.quantidadeRel = qtd
            self.categoriaRel = cate

            paragrafo = doc.add_paragraph()





            add_codigo_negrito= paragrafo.add_run('Código:')
            add_codigo_negrito.bold = True
            add_codigo_negrito.font.name = 'Arial'
            add_codigo_negrito.font.size = Pt(12)
            add_codigo=paragrafo.add_run(f' {self.codigoRel}\n')
            add_codigo.font.name='Arial'
            add_codigo.font.size = Pt(12)


            add_produto_negrito = paragrafo.add_run('Produto:')
            add_produto_negrito.bold = True
            add_produto_negrito.font.name = 'Arial'
            add_produto_negrito.font.size = Pt(12)
            add_produto = paragrafo.add_run(f' {self.produtoRel}\n')
            add_produto.font.name = 'Arial'
            add_produto.font.size = Pt(12)

            add_qtd_negrito = paragrafo.add_run('Quantidade:')
            add_qtd_negrito.bold = True
            add_qtd_negrito.font.name = 'Arial'
            add_qtd_negrito.font.size = Pt(12)
            add_qtd = paragrafo.add_run(f' {self.quantidadeRel}\n')
            add_qtd.font.name = 'Arial'
            add_qtd.font.size = Pt(12)


            add_valor_negrito = paragrafo.add_run('Valor:')
            add_valor_negrito.bold = True
            add_valor_negrito.font.name = 'Arial'
            add_valor_negrito.font.size = Pt(12)
            add_valor = paragrafo.add_run(f' R$ {self.categoriaRel}\n')
            add_valor.font.name = 'Arial'
            add_valor.font.size = Pt(12)



            add_linha = paragrafo.add_run('------------------------------------------------------------------------------------------------------------')

        hoje= datetime.now().strftime("%d/%m/%Y")
        hj=hoje.replace('/','-')

        doc.save(f'Estoque do dia {hj}.docx')
        cursor.close()
        banco.close()
        msg = 'Relatório gerado com Sucesso!'
        messagebox.showinfo('Aviso!', msg)


class Funcoes():
    def limpa_tela(self):
        self.codigo_entry.delete(0, END)
        self.nomeProduto.delete(0, END)
        self.quantProduto.delete(0, END)
        self.categoriaProduto.delete(0, END)

    def conecta_bd(self):
        self.co = sqlite3.connect('produtos.bd')
        self.cursor = self.co.cursor();


    def desconecta_bd(self):
        self.co.close()

    def montaTabelas(self):
        self.conecta_bd()

        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS produtos(
                cod INTEGER PRIMARY KEY, 
                nome_produto CHAR(40) NOT NULL, 
                quantidade INTEGER(20) NOT NULL, 
                categoria CHAR(30)
            ); 
        ''')

        self.co.commit();

        self.desconecta_bd()

    def variaveis(self):
        self.codigo = self.codigo_entry.get()
        self.nome = self.nomeProduto.get()
        self.quantidade = self.quantProduto.get()
        self.categoria = self.categoriaProduto.get()

    def add_produto(self):
        self.variaveis()

        if self.nomeProduto.get() == '' and self.quantProduto.get() == '' and self.categoriaProduto.get() == '':
            msg = 'Os campos estão vazios!'
            messagebox.showinfo('Aviso!', msg)

        elif self.nomeProduto.get() == '':
            msg = 'Informar o NOME dos produtos!'
            messagebox.showinfo('Aviso!', msg)

        elif self.quantProduto.get() == '':
            msg = 'Informar a QUANTIDADE de produtos!'
            messagebox.showinfo('Aviso!', msg)

        elif self.categoriaProduto.get() == '':
            msg = 'Informar o VALOR dos produtos!'
            messagebox.showinfo('Aviso!', msg)
        else:
            self.conecta_bd()
            self.cursor.execute(''' INSERT INTO produtos (nome_produto, quantidade, categoria)
                VALUES (?, ?, ?)''', (self.nome, self.quantidade, self.categoria))
            self.co.commit()
            self.desconecta_bd()
            self.select_lista()
            self.limpa_tela()

    def select_lista(self):
        self.listaCli.delete(*self.listaCli.get_children())
        self.conecta_bd()
        lista = self.cursor.execute(''' SELECT cod, nome_produto, quantidade, categoria FROM produtos
            ORDER BY nome_produto ASC; ''')

        for i in lista:

            self.listaCli.insert("", END, values=i)

        self.desconecta_bd()

    def OnDoubleClick(self, event):
        self.limpa_tela()
        self.listaCli.selection()

        for n in self.listaCli.selection():
            col1, col2, col3, col4 = self.listaCli.item(n, 'values')
            self.codigo_entry.insert(END, col1)
            self.nomeProduto.insert(END, col2)
            self.quantProduto.insert(END, col3)
            self.categoriaProduto.insert(END, col4)

    def deleta_produto(self):
        self.variaveis()
        self.conecta_bd()
        self.cursor.execute("""DELETE FROM produtos WHERE cod = ?""", (self.codigo,))
        self.co.commit()
        self.desconecta_bd()
        self.limpa_tela()
        self.select_lista()

    def altera_produto(self):
        self.variaveis()
        self.conecta_bd()
        self.cursor.execute(""" Update produtos SET nome_produto = ?, quantidade = ?, categoria = ? WHERE cod = ?""",
                            (self.nome, self.quantidade, self.categoria, self.codigo))
        self.co.commit()
        self.desconecta_bd()
        self.select_lista()
        self.limpa_tela()

    def busca_produto(self):
        self.conecta_bd()
        self.listaCli.delete(*self.listaCli.get_children())

        self.nomeProduto.insert(END, '%')
        nome = self.nomeProduto.get()
        self.cursor.execute(
            """SELECT cod , nome_produto, quantidade , categoria FROM produtos WHERE nome_produto LIKE '%s' ORDER BY nome_produto ASC""" % nome)
        buscanomeCLI = self.cursor.fetchall()
        for i in buscanomeCLI:
            self.listaCli.insert("", END, values=i)
        self.limpa_tela()
        self.desconecta_bd()

class Application(Funcoes, Relatorios):
    def __init__(self):
        self.janela = janela
        self.tela()
        self.frames_tela()
        self.widgets_frame1()
        self.lista_frame2()
        self.montaTabelas()
        self.select_lista()
        self.Menus()
        janela.mainloop()

    def tela(self):
        self.janela.title('2 IRMÃOS ESTOQUE')
        self.janela.configure(background='white')
        self.janela.geometry('788x588')
        self.janela.resizable(True, True)
        self.janela.minsize(width=788, height=588)

    def frames_tela(self):
        self.frame1 = Frame(self.janela, bd=4, bg='#dfe3ee', highlightbackground='black', highlightthickness=3)
        self.frame1.place(relx=0.02, rely=0.02, relwidth=0.96, relheight=0.46)

        self.frame2 = Frame(self.janela, bd=4, bg='#dfe3ee', highlightbackground='black', highlightthickness=3)
        self.frame2.place(relx=0.02, rely=0.5, relwidth=0.96, relheight=0.45)

    def widgets_frame1(self):
        # botão limpar...
        azul='#3498DB'
        preto='black'
        self.bt_limpar = tk.Button(self.frame1, text='Limpar',  bg= azul, fg= 'black' ,font=("verdana", 10,'bold'),
                                command=self.limpa_tela)
        self.bt_limpar.place(relx=0.2, rely=0.1, relwidth=0.1, relheight=0.15)

        # botão buscar...
        self.bt_buscar = tk.Button(self.frame1, text='Buscar', bg= azul, fg= preto ,font=('verdana', 10,'bold'), command = self.busca_produto)
        self.bt_buscar.place(relx=0.31, rely=0.1, relwidth=0.1, relheight=0.15)

        # botão novo...
        self.bt_novo = tk.Button(self.frame1, text='Novo', bg= azul, fg= preto ,font=('verdana',10,'bold'),command=self.add_produto)
        self.bt_novo.place(relx=0.59, rely=0.1, relwidth=0.1, relheight=0.15)

        # botão alterar...
        self.bt_alterar = tk.Button(self.frame1, text='Alterar', bg= azul, fg= preto ,font=('verdana', 10,'bold'),command=self.altera_produto)
        self.bt_alterar.place(relx=0.7, rely=0.1, relwidth=0.1, relheight=0.15)

        # botão apagar...
        self.bt_apagar = tk.Button(self.frame1, text='Apagar', bg= azul, fg= preto ,font=('verdana', 10,'bold'),
                                command=self.deleta_produto)
        self.bt_apagar.place(relx=0.81, rely=0.1, relwidth=0.1, relheight=0.15)

        # criação label e entrada do codigo...
        self.label_codigo = Label(self.frame1, text='Código', bg='#dfe3ee', fg=preto)
        self.label_codigo.place(relx=0.05, rely=0.05)

        self.codigo_entry = Entry(self.frame1)
        self.codigo_entry.place(relx=0.05, rely=0.15, relwidth=0.06)

        # criação label e entrada do nome do produto...
        self.label_nomeProduto = Label(self.frame1, text='Nome do Produto', bg='#dfe3ee', fg=preto)
        self.label_nomeProduto.place(relx=0.05, rely=0.35)

        self.nomeProduto = Entry(self.frame1)
        self.nomeProduto.place(relx=0.05, rely=0.45, relwidth=0.85)

        # criação label e entrada da quantidade de produtos...
        self.label_quantProduto = Label(self.frame1, text='Quantidade', bg='#dfe3ee', fg=preto)
        self.label_quantProduto.place(relx=0.05, rely=0.6)

        self.quantProduto = Entry(self.frame1)
        self.quantProduto.place(relx=0.05, rely=0.7, relwidth=0.4)

        # criação label e entrada da categoria do produto...
        self.label_categoriaProduto = Label(self.frame1, text='Valor', bg='#dfe3ee', fg=preto)
        self.label_categoriaProduto.place(relx=0.5, rely=0.6)

        self.categoriaProduto = Entry(self.frame1)
        self.categoriaProduto.place(relx=0.5, rely=0.7, relwidth=0.4)

    def lista_frame2(self):
        self.listaCli = ttk.Treeview(self.frame2, height=3, column=("col1", "col2", "col3", "col4"))
        self.listaCli.column('#0', width=1, stretch=NO)
        self.listaCli.heading("#1", text='Código')
        self.listaCli.heading("#2", text='Nome do produto')
        self.listaCli.heading("#3", text='Quantidade')
        self.listaCli.heading("#4", text='Valor')

        self.listaCli.column("#0", width=1)
        self.listaCli.column("#1", width=50)
        self.listaCli.column("#2", width=200)
        self.listaCli.column("#3", width=125)
        self.listaCli.column("#4", width=125)
        style = ttk.Style(janela)
        style.theme_use('clam')

        self.listaCli.place(relx=0.01, rely=0.1, relwidth=0.95, relheight=0.85)

        self.scroolLista = Scrollbar(self.frame2, orient='vertical',
                                     command=self.listaCli.yview)
        self.listaCli.configure(yscroll=self.scroolLista.set)
        self.scroolLista.place(relx=0.96, rely=0.1, relwidth=0.04, relheight=0.85)
        self.listaCli.bind("<Double-1>", self.OnDoubleClick)

    def Menus(self):
        menubar = Menu(self.janela)
        self.janela.config(menu=menubar)
        filemenu = Menu(menubar, tearoff=0)
        filemenu2 = Menu(menubar, tearoff=0)

        def Quit(): self.janela.destroy()

        menubar.add_cascade(label="Opções", menu=filemenu)
        menubar.add_cascade(label="Relatório", menu=filemenu2)

        filemenu.add_command(label="Sair", command=Quit)
        filemenu2.add_command(label="Gerar Relatório", command=self.geraRelatorioProduto)

Application()